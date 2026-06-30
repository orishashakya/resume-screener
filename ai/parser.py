"""
ai/parser.py
Phase 1: Resume Parsing
Extracts clean text from PDF and DOCX resume files.
"""

import pdfplumber
import docx
import re
import os


def extract_text_from_pdf(filepath):
    """Extract raw text from a PDF file, page by page."""
    text_parts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(filepath):
    """Extract raw text from a DOCX file, including tables."""
    doc = docx.Document(filepath)
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


def clean_text(text):
    """Normalize whitespace and strip common PDF extraction artifacts."""
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"●|•|▪", "-", text)
    return text.strip()


def parse_resume(filepath):
    """
    Main entry point. Takes a path to a .pdf or .docx resume
    and returns cleaned plain text.
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        raw = extract_text_from_pdf(filepath)
    elif ext == ".docx":
        raw = extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use .pdf or .docx")

    if not raw or not raw.strip():
        raise ValueError(f"No extractable text found in {filepath} "
                          f"(it may be a scanned/image-based PDF).")

    return clean_text(raw)


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        result = parse_resume(sys.argv[1])
        print(result[:1000])
        print(f"\n--- Total characters extracted: {len(result)} ---")
    else:
        print("Usage: python parser.py <path_to_resume.pdf|docx>")